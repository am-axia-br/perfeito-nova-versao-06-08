# ===================================================================
# app/exportador_docx.py (VERSÃO FINAL COM CONTEÚDO IGUAL À INTERFACE)
#
# O que mudou:
# - PRESERVAÇÃO DE CONTEÚDO: O documento Word agora mostra exatamente 
#   o mesmo conteúdo que é exibido na interface, mantendo todas as 
#   informações visíveis para o usuário
# - FORMATO VISUAL APRIMORADO: Mantém o layout visual solicitado com
#   cabeçalhos coloridos, tabelas estruturadas e blocos PJe-Calc
# - MÁXIMA FIDELIDADE: Garante que o documento impresso reflita
#   fielmente o que o usuário vê na tela
# ===================================================================

from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
import os
import json
import re

def format_value(value):
    """Converte de forma inteligente qualquer valor para uma string legível."""
    if value is None or value == '':
        return 'Não informado'
    if isinstance(value, str):
        return value
    if isinstance(value, list):
        return ', '.join([format_value(item) for item in value])
    if isinstance(value, dict):
        return ', '.join([f"{str(k).replace('_', ' ').title()}: {v}" for k, v in value.items()])
    return str(value)

def rgb_from_hex(hex_str):
    """Converte cor hexadecimal para RGB."""
    hex_str = hex_str.lstrip('#')
    return RGBColor(int(hex_str[0:2], 16), int(hex_str[2:4], 16), int(hex_str[4:6], 16))

def gerar_docx_resumo(dados, caminho_saida="export/resumo_processo.docx"):
    """Gera um resumo do processo em formato .docx usando exatamente os dados exibidos na interface."""
    doc = Document()
    
    # Configurações de página
    sections = doc.sections
    for section in sections:
        section.page_height = Inches(11)
        section.page_width = Inches(8.5)
        section.left_margin = Inches(1)
        section.right_margin = Inches(1)
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
    
    # Função auxiliar para adicionar cabeçalhos coloridos
    def adicionar_cabecalho(texto, cor="#ff0000"):
        p = doc.add_paragraph()
        run = p.add_run(f"✓ {texto}")
        run.bold = True
        run.font.size = Pt(14)
        run.font.color.rgb = rgb_from_hex(cor)
        return p
    
    # 1. CABEÇALHO - RESUMO TÉCNICO DO PROCESSO
    adicionar_cabecalho("Resumo Técnico do Processo")
    
    # 2. OBTER RESUMO DO PROCESSO EXATAMENTE COMO ESTÁ NA INTERFACE
    resumo = dados.get("observacoes_gerais", "")
    
    # Limpar prefixos indesejados, se houver
    if "✅ **Resumo Técnico do Processo**" in resumo:
        resumo_limpo = resumo.replace("✅ **Resumo Técnico do Processo** - ", "")
    else:
        resumo_limpo = resumo
    
    # Adicionar resumo ao documento
    p = doc.add_paragraph()
    p.add_run(resumo_limpo)
    
    # 3. VERBAS PLEITEADAS
    doc.add_paragraph()
    adicionar_cabecalho("Verbas Pleiteadas (petição inicial) e Parâmetros Relevantes", "#4CAF50")
    
    # Extrair informações de verbas do texto formatado, se disponíveis
    verbas = []
    verbas_texto = ""
    
    # Procurar bloco de verbas pleiteadas na interface
    match = re.search(r"✅ Verbas Pleiteadas \(Petição Inicial\)(.*?)(?:✅|$)", resumo, re.DOTALL)
    if match:
        verbas_texto = match.group(1).strip()
        
        # Adicionar texto original de verbas
        p = doc.add_paragraph()
        p.add_run(verbas_texto)
        
        # Tentar extrair as verbas para a tabela
        categorias = re.findall(r'([^:]+):(.*?)(?=\n[A-Z]|\Z)', verbas_texto, re.DOTALL)
        for categoria, detalhes in categorias:
            for item in detalhes.strip().split(', '):
                verbas.append({
                    "verba": f"{categoria.strip()}: {item.strip()}",
                    "periodo": "",
                    "reflexos": ""
                })
    else:
        # Usar verbas do dicionário de dados se não encontrar no texto
        verbas = dados.get('verbas_pleiteadas', [])
        if not verbas:
            pleitos = dados.get("pleitos_e_verbas", [])
            for pleito in pleitos:
                verbas.append({
                    "verba": format_value(pleito.get("verba", "")),
                    "periodo": format_value(pleito.get("parametros", "")),
                    "reflexos": ""
                })
    
    # 4. ATUALIZAÇÃO MONETÁRIA
    doc.add_paragraph()
    adicionar_cabecalho("ATUALIZAÇÃO MONETÁRIA", "#9C27B0")
    
    # Procurar bloco de atualização monetária na interface
    monetaria_texto = ""
    match = re.search(r"✅ Atualização Monetária e Parâmetros(.*?)(?:✅|$)", resumo, re.DOTALL)
    if match:
        monetaria_texto = match.group(1).strip()
        
        # Adicionar texto original de atualização monetária
        p = doc.add_paragraph()
        p.add_run(monetaria_texto)
    else:
        # Tabela de atualização monetária com dados do dicionário
        tabela_monetaria = doc.add_table(rows=3, cols=2)
        tabela_monetaria.style = 'Table Grid'
        
        # Preencher tabela monetária
        cells = tabela_monetaria.rows[0].cells
        cells[0].text = "Campo"
        cells[1].text = "Valor"
        
        info_pjecalc = dados.get("informacoes_pjecalc", {})
        
        cells = tabela_monetaria.rows[1].cells
        cells[0].text = "Índice monetário"
        cells[1].text = format_value(info_pjecalc.get("correcao_monetaria", "IPCA-E"))
        
        cells = tabela_monetaria.rows[2].cells
        cells[0].text = "Juros de mora"
        cells[1].text = format_value(info_pjecalc.get("juros_mora", "TRD pré judicial"))
    
    # 5. PREMISSA ESPECÍFICA
    doc.add_paragraph()
    adicionar_cabecalho("PREMISSA ESPECÍFICA CARTEIRA:", "#FF5722")
    p = doc.add_paragraph("Inidoxy")
    p = doc.add_paragraph()
    p.add_run("INSS - Terceiros: ").bold = True
    p.add_run("5,8% (lançar manualmente)")
    
    # 6. REVISÃO FINAL DO QUADRO PARA O CÁLCULO
    doc.add_paragraph()
    adicionar_cabecalho("REVISÃO FINAL DO QUADRO PARA O CÁLCULO", "#2196F3")
    
    # Procurar quadro para cálculo na interface
    quadro_texto = ""
    match = re.search(r"✅ Revisão Final – Quadro para Cálculo(.*?)(?:✅|$)", resumo, re.DOTALL)
    if match:
        quadro_texto = match.group(1).strip()
        
        # Tentar extrair as linhas da tabela
        linhas_tabela = re.findall(r'([^\t\n]+)\t([^\t\n]+)\t([^\t\n]*)', quadro_texto)
        
        if linhas_tabela:
            tabela_final = doc.add_table(rows=1, cols=3)
            tabela_final.style = 'Table Grid'
            
            # Cabeçalhos
            headers = tabela_final.rows[0].cells
            headers[0].text = "Verba"
            headers[1].text = "Parâmetro"
            headers[2].text = "Reflexos"
            
            # Formatar cabeçalhos
            for cell in headers:
                cell.paragraphs[0].runs[0].bold = True
            
            # Adicionar linhas
            for verba, parametro, reflexos in linhas_tabela:
                row_cells = tabela_final.add_row().cells
                row_cells[0].text = verba.strip()
                row_cells[1].text = parametro.strip()
                row_cells[2].text = reflexos.strip()
        else:
            # Adicionar texto original se não conseguir extrair a tabela
            p = doc.add_paragraph()
            p.add_run(quadro_texto)
    else:
        # Usar dados do dicionário se não encontrar no texto
        tabela_final = doc.add_table(rows=1, cols=3)
        tabela_final.style = 'Table Grid'
        
        # Cabeçalhos
        headers = tabela_final.rows[0].cells
        headers[0].text = "Verba"
        headers[1].text = "Configurações PJE-CALC"
        headers[2].text = "Reflexos"
        
        # Formatar cabeçalhos
        for cell in headers:
            cell.paragraphs[0].runs[0].bold = True
        
        # Adicionar verbas para revisão final
        for verba in verbas:
            row_cells = tabela_final.add_row().cells
            row_cells[0].text = format_value(verba.get('verba', ''))
            row_cells[1].text = format_value(verba.get('periodo', '') or verba.get('configuracao_pjecalc', ''))
            row_cells[2].text = format_value(verba.get('reflexos', ''))
    
    # 7. MODELO - LAUDO DE INSTRUÇÃO TÉCNICA
    doc.add_paragraph()
    adicionar_cabecalho("MODELO - LAUDO DE INSTRUÇÃO TÉCNICA PARA O CÁLCULO (PJe-Calc)", "#795548")
    
    # Procurar modelo laudo técnico na interface
    laudo_texto = ""
    match = re.search(r"✅ Modelo Laudo Técnico para PJe-Calc \(com BLOCOS 1 a 5\)(.*?)(?:Processado por|$)", resumo, re.DOTALL)
    if match:
        laudo_texto = match.group(1).strip()
        
        # Adicionar texto original do laudo técnico
        p = doc.add_paragraph()
        p.add_run(laudo_texto)
    else:
        # Blocos padrão se não encontrar no texto
        # BLOCO 1
        p = doc.add_paragraph()
        p.add_run("BLOCO 1 - DADOS CADASTRAIS").bold = True
        
        dados_pessoais = dados.get("dados_pessoais", {})
        info_pjecalc = dados.get("informacoes_pjecalc", {})
        
        # Usar os dados do dicionário
        reclamante = format_value(dados_pessoais.get("reclamante", ""))
        reclamada = format_value(dados_pessoais.get("reclamada", ""))
        
        # Dividir reclamadas se houver múltiplas
        reclamadas_separadas = []
        if ';' in reclamada:
            reclamadas_separadas = [r.strip() for r in reclamada.split(';')]
        elif ',' in reclamada:
            reclamadas_separadas = [r.strip() for r in reclamada.split(',')]
        else:
            reclamadas_separadas = [reclamada]
        
        # Garantir que temos pelo menos duas reclamadas (mesmo que vazias)
        if len(reclamadas_separadas) < 2:
            reclamadas_separadas.append('')
        
        dados_cadastrais = [
            f"Processo: {format_value(info_pjecalc.get('numero_processo', ''))}",
            f"Reclamante: {reclamante}",
            f"CPF: {format_value(dados_pessoais.get('cpf', ''))}",
            f"Reclamada 1: {reclamadas_separadas[0]}",
            f"Reclamada 2: {reclamadas_separadas[1] if len(reclamadas_separadas) > 1 else ''}",
            f"Data de Ajuizamento: {format_value(info_pjecalc.get('data_liquidacao', ''))}",
            f"Período Contratual: {format_value(dados_pessoais.get('data_admissao', ''))} a {format_value(dados_pessoais.get('data_demissao', ''))} - {format_value(dados_pessoais.get('tipo_rescisao', 'rescisão indireta'))}",
            f"Cidade: {format_value(info_pjecalc.get('jurisdicao', '')).split('/')[0] if '/' in format_value(info_pjecalc.get('jurisdicao', '')) else ''}",
            f"Períodos Afastamentos Citados: {format_value(dados_pessoais.get('periodo_afastamento', ''))}",
            f"Salário Base Contratual: {format_value(dados_pessoais.get('ultimo_salario', ''))}",
            f"Valor da Causa: {format_value(info_pjecalc.get('valor_causa', ''))}",
            f"Fase Cálculo: {format_value(info_pjecalc.get('fase_calculo', 'Provisão Inicial'))}"
        ]
        
        for item in dados_cadastrais:
            lista = doc.add_paragraph(style='List Bullet')
            lista.add_run(item)
        
        # BLOCO 2
        p = doc.add_paragraph()
        p.add_run("BLOCO 2 - CADASTRO AFASTAMENTOS / HISTÓRICO SALARIAL").bold = True
        
        lista = doc.add_paragraph(style='List Bullet')
        lista.add_run(f"Salário Base Contratual: {format_value(dados_pessoais.get('ultimo_salario', ''))}")
        
        # BLOCO 3, 4 e 5 (conforme documentação anterior)
        p = doc.add_paragraph()
        p.add_run("BLOCO 3 - VERBAS A APURAR").bold = True
        p = doc.add_paragraph()
        p.add_run("BLOCO 4 - HONORÁRIOS / CUSTAS / INSS E CORREÇÃO").bold = True
        p = doc.add_paragraph()
        p.add_run("BLOCO 5 - LIQUIDAÇÃO E IMPRESSÃO").bold = True
    
    # 8. SEÇÃO FINAL - DADOS PESSOAIS E INFORMAÇÕES PJE-CALC
    doc.add_paragraph()
    
    # Procurar dados pessoais e contratuais na interface
    match_pessoais = re.search(r"📌 Dados Pessoais e Contratuais(.*?)(?:🗂️|$)", resumo, re.DOTALL)
    if match_pessoais:
        dados_pessoais_texto = match_pessoais.group(1).strip()
        
        p = doc.add_paragraph()
        p.add_run("📌 Dados Pessoais e Contratuais").bold = True
        p = doc.add_paragraph()
        p.add_run(dados_pessoais_texto)
    
    # Procurar informações para preenchimento na interface
    match_info = re.search(r"🗂️ Informações para Preenchimento no PJe-Calc(.*?)(?:📑|$)", resumo, re.DOTALL)
    if match_info:
        info_texto = match_info.group(1).strip()
        
        doc.add_paragraph()
        p = doc.add_paragraph()
        p.add_run("🗂️ Informações para Preenchimento no PJe-Calc").bold = True
        p = doc.add_paragraph()
        p.add_run(info_texto)
    
    # 9. RESULTADO FINAL DA LIQUIDAÇÃO
    match_resultado = re.search(r"🟢 Resultado Final da Liquidação(.*?)(?:📥|$)", resumo, re.DOTALL)
    if match_resultado:
        resultado_texto = match_resultado.group(1).strip()
        
        doc.add_paragraph()
        p = doc.add_paragraph()
        p.add_run("🟢 Resultado Final da Liquidação").bold = True
        p = doc.add_paragraph()
        p.add_run(resultado_texto)
    
    # 10. Adicionar data de processamento no rodapé
    match_processado = re.search(r"Processado por (.*?) em (.*?) UTC", resumo)
    if match_processado:
        usuario = match_processado.group(1)
        data = match_processado.group(2)
        
        doc.add_paragraph()
        doc.add_paragraph(f"Processado por {usuario} em {data} UTC.")
    else:
        # Usar data e usuário fornecidos
        doc.add_paragraph()
        doc.add_paragraph(f"Processado por am-axia-br em 2025-08-06 02:28:10 UTC.")
    
    # Salvar documento
    os.makedirs(os.path.dirname(caminho_saida), exist_ok=True)
    doc.save(caminho_saida)
    
    return caminho_saida